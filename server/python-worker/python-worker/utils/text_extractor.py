import re
import os
import logging
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file using PyPDF2.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted raw text from the PDF
        
    Raises:
        Exception: If PDF reading fails
    """
    try:
        logger.info(f"📄 Extracting text from PDF: {file_path}")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        reader = PdfReader(file_path)
        text = ""
        
        # Extract text from all pages
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
                logger.debug(f"  - Extracted page {page_num}: {len(page_text)} characters")
        
        if not text.strip():
            logger.warning("⚠️  PDF appears to be empty or text extraction failed")
            return ""
        
        logger.info(f"✅ Successfully extracted {len(text)} characters from PDF")
        return text
        
    except Exception as error:
        logger.error(f"❌ Error extracting text from PDF: {error}")
        raise


def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted raw text from the DOCX
        
    Raises:
        Exception: If DOCX reading fails
    """
    try:
        logger.info(f"📄 Extracting text from DOCX: {file_path}")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from all paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            logger.warning("⚠️  DOCX appears to be empty")
            return ""
        
        logger.info(f"✅ Successfully extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as error:
        logger.error(f"❌ Error extracting text from DOCX: {error}")
        raise


def extract_text_from_doc(file_path):
    """
    Extract text from a DOC file (legacy Microsoft Word format).
    This is a STUB implementation. Full support requires additional libraries like textract.
    
    Args:
        file_path (str): Path to the DOC file
        
    Returns:
        str: Empty string (stub implementation)
        
    Raises:
        NotImplementedError: DOC format not yet supported
    """
    logger.warning(f"⚠️  DOC format not fully supported yet: {file_path}")
    logger.info("💡 Suggestion: Please convert DOC to DOCX or PDF for text extraction")
    
    # Stub implementation - to be enhanced in future
    raise NotImplementedError(
        "DOC format extraction not implemented. "
        "Please use PDF or DOCX formats."
    )


def clean_text(raw_text):
    """
    Clean and normalize extracted text.
    - Remove excessive whitespace
    - Normalize line breaks
    - Remove special characters (optional)
    - Trim leading/trailing spaces
    
    Args:
        raw_text (str): Raw extracted text
        
    Returns:
        str: Cleaned and normalized text
    """
    if not raw_text:
        return ""
    
    logger.debug("🧹 Cleaning extracted text...")
    
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', raw_text)
    
    # Replace multiple newlines with double newline (paragraph separator)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove leading/trailing whitespace from entire text
    text = text.strip()
    
    # Remove null characters and other control characters
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
    
    logger.debug(f"✅ Text cleaned: {len(text)} characters")
    return text


def extract_text(file_path):
    """
    Main extraction function that detects file type and extracts text accordingly.
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Cleaned extracted text
        
    Raises:
        ValueError: If file extension is not supported
        Exception: If extraction fails
    """
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    logger.info(f"🔍 Detected file extension: {ext}")
    
    # Extract based on file type
    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        raw_text = extract_text_from_docx(file_path)
    elif ext == '.doc':
        raw_text = extract_text_from_doc(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .pdf, .docx")
    
    # Clean the extracted text
    cleaned_text = clean_text(raw_text)
    
    return cleaned_text
